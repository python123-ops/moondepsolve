from __future__ import annotations

from pathlib import Path
import hashlib

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfdoc, pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "competition"
PDF_PATH = OUT_DIR / "MoonDepSolve项目申报书.pdf"
DOCX_PATH = OUT_DIR / "MoonDepSolve项目申报书.docx"

TITLE = "MoonDepSolve v0.3 项目申报书"
SUBTITLE = "MoonBit 包生态语义版本、依赖求解与升级规划基础库"
GITLINK_LABEL = "gitlink.org.cn/python123/moondepsolve"
GITHUB_LABEL = "github.com/python123-ops/moondepsolve"
AUTHOR = "python123"

SECTIONS = [
    (
        "项目定位",
        "MoonDepSolve 面向 MoonBit 包生态，覆盖版本约束、传递依赖求解、稳定 lock、依赖图、冲突解释和升级规划，可复用于包管理器、构建工具、依赖审计与自动化发布。",
    ),
    (
        "v0.3 核心成果",
        "保持 v0.1/v0.2 API 兼容；新增 HighestCompatible 与精确 MinimalChange，先最小化变更包数量，同成本时稳定偏好更高版本；新增 UpgradePlan 及有界搜索错误。原生文件 CLI 可从 registry/lock 执行 resolve 和 plan，输出 lock、文本图或 Graphviz DOT。",
    ),
    (
        "技术路线与生态价值",
        "采用“版本解析 -> 约束匹配 -> 候选稳定排序 -> 递归求解/有界精确搜索 -> lock/图/诊断输出”边界。核心算法使用 MoonBit，公共接口由 moon info 生成的 .mbti 审核；文件工具使用官方 moonbitlang/async/fs，为 MoonBit 工具链补充可复用、可解释、可测试的依赖能力。",
    ),
    (
        "工程质量与公开维护",
        "默认后端 27 项、native 后端 31 项测试通过，另有 4 组 CLI expected 输出回归。CI 拉取完整历史，先检查 author/committer 仅为 python123，再检查接口、格式、诊断、覆盖率和 native CLI。仓库提供 Apache-2.0、README、Changelog、贡献/安全规范、模板与第三方许可证记录。",
    ),
    (
        "赛事进度与交付",
        "官方开发期为 2026-04-29 至 2026-07-12，验收期为 2026-07-13 至 2026-07-17。4 月 29 日后完成 v0.1 求解基础、v0.2 图与冲突解释、v0.3 精确升级规划和文件 CLI。终验交付双仓库一致历史、v0.3.0 标签/Release、fresh clone 复验与 Mooncakes dry-run；发布由 python123 授权执行。",
    ),
]


def _compatible_md5(data=b"", **_kwargs):
    return hashlib.md5(data)


pdfdoc.md5 = _compatible_md5


def register_font() -> str:
    for path in [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]:
        if path.exists():
            pdfmetrics.registerFont(TTFont("MoonDepSolveCN", str(path)))
            return "MoonDepSolveCN"
    return "Helvetica"


def build_pdf() -> None:
    font = register_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCN",
        parent=styles["Title"],
        fontName=font,
        fontSize=23,
        leading=28,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#17365D"),
        spaceAfter=3,
    )
    subtitle_style = ParagraphStyle(
        "SubtitleCN",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=11.5,
        leading=15,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#4B5563"),
        spaceAfter=14,
    )
    heading_style = ParagraphStyle(
        "HeadingCN",
        parent=styles["Heading2"],
        fontName=font,
        fontSize=13,
        leading=17,
        textColor=colors.HexColor("#1F4D78"),
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "BodyCN",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=10.2,
        leading=16,
        firstLineIndent=20.4,
        textColor=colors.HexColor("#222222"),
        spaceAfter=3,
    )
    meta_style = ParagraphStyle(
        "MetaCN",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=8.7,
        leading=11.5,
        textColor=colors.HexColor("#263238"),
    )

    document = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=1.25 * cm,
        rightMargin=1.25 * cm,
        topMargin=1.05 * cm,
        bottomMargin=1.05 * cm,
        title=TITLE,
        author=AUTHOR,
        subject=SUBTITLE,
    )
    story = [
        Paragraph(TITLE, title_style),
        Paragraph(SUBTITLE, subtitle_style),
    ]
    metadata = [
        [Paragraph("<b>作者</b>", meta_style), Paragraph(AUTHOR, meta_style),
         Paragraph("<b>许可证</b>", meta_style), Paragraph("Apache-2.0", meta_style)],
        [Paragraph("<b>GitLink</b>", meta_style), Paragraph(GITLINK_LABEL, meta_style),
         Paragraph("<b>GitHub</b>", meta_style), Paragraph(GITHUB_LABEL, meta_style)],
    ]
    table = Table(
        metadata,
        colWidths=[1.6 * cm, 6.8 * cm, 1.6 * cm, 8.1 * cm],
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8EEF5")),
                ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#E8EEF5")),
                ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#9AA9B8")),
                ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#C5CDD5")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    story.extend([table, Spacer(1, 0.15 * cm)])
    for heading, text in SECTIONS:
        story.append(Paragraph(heading, heading_style))
        story.append(Paragraph(text, body_style))
    document.build(story)

    pages = len(PdfReader(str(PDF_PATH)).pages)
    if pages != 1:
        raise RuntimeError(f"Proposal PDF must be exactly one page, got {pages}")


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_width(cell, width_twips: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_twips))
    width.set(qn("w:type"), "dxa")


def set_cell_margins(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")

    table_indent = properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:w"), str(indent))
    table_indent.set(qn("w:type"), "dxa")

    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def set_run_font(run, name: str, size: float, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.05)
    section.header_distance = Cm(0.4)
    section.footer_distance = Cm(0.4)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.15
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    heading_style = document.styles["Heading 1"]
    heading_style.font.name = "Microsoft YaHei"
    heading_style.font.size = Pt(13)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
    heading_style.paragraph_format.space_before = Pt(10)
    heading_style.paragraph_format.space_after = Pt(4)
    heading_style.paragraph_format.keep_with_next = True
    heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run(TITLE)
    set_run_font(run, "Microsoft YaHei", 23, True)
    run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    run = subtitle.add_run(SUBTITLE)
    set_run_font(run, "Microsoft YaHei", 11.5)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    table = document.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [950, 3750, 950, 4690]
    set_table_geometry(table, widths)
    rows = [
        ("作者", AUTHOR, "许可证", "Apache-2.0"),
        ("GitLink", GITLINK_LABEL, "GitHub", GITHUB_LABEL),
    ]
    for row, values in zip(table.rows, rows):
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, "Microsoft YaHei", 8.7, index in (0, 2))
            if index in (0, 2):
                set_cell_shading(cell, "E8EEF5")

    for heading, text in SECTIONS:
        document.add_paragraph(heading, style="Heading 1")

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(20.4)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(text)
        set_run_font(run, "Microsoft YaHei", 10.2)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("MoonDepSolve v0.3 | 2026-06-21")
    set_run_font(run, "Microsoft YaHei", 7.5)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    document.core_properties.title = TITLE
    document.core_properties.subject = SUBTITLE
    document.core_properties.author = AUTHOR
    document.core_properties.last_modified_by = AUTHOR
    document.core_properties.keywords = "MoonBit, dependency resolver, OSC2026"
    document.save(DOCX_PATH)


def validate_docx() -> None:
    document = Document(DOCX_PATH)
    if len(document.sections) != 1:
        raise RuntimeError("Proposal DOCX must contain exactly one section")

    section = document.sections[0]
    expected_dimensions = {
        "page width": (section.page_width, Cm(21)),
        "page height": (section.page_height, Cm(29.7)),
        "left margin": (section.left_margin, Cm(1.25)),
        "right margin": (section.right_margin, Cm(1.25)),
        "top margin": (section.top_margin, Cm(1.05)),
        "bottom margin": (section.bottom_margin, Cm(1.05)),
    }
    for label, (actual, expected) in expected_dimensions.items():
        # Word stores section dimensions in twips, so round-trip values differ
        # from exact centimetre conversions by up to half a twip.
        if abs(actual - expected) > 4000:
            raise RuntimeError(f"Unexpected DOCX {label}: {actual}")

    properties = document.core_properties
    if properties.title != TITLE or properties.subject != SUBTITLE:
        raise RuntimeError("DOCX title metadata is out of date")
    if properties.author != AUTHOR or properties.last_modified_by != AUTHOR:
        raise RuntimeError("DOCX contributor metadata must be python123")

    headings = [p for p in document.paragraphs if p.style.name == "Heading 1"]
    if [p.text for p in headings] != [heading for heading, _ in SECTIONS]:
        raise RuntimeError("DOCX Heading 1 structure does not match proposal sections")

    if len(document.tables) != 1:
        raise RuntimeError("Proposal DOCX must contain exactly one metadata table")
    table = document.tables[0]
    expected_widths = [950, 3750, 950, 4690]
    grid_widths = [
        int(column.get(qn("w:w"))) for column in table._tbl.tblGrid.gridCol_lst
    ]
    if grid_widths != expected_widths:
        raise RuntimeError(f"Unexpected DOCX table grid: {grid_widths}")
    for row in table.rows:
        widths = [int(cell._tc.tcPr.tcW.get(qn("w:w"))) for cell in row.cells]
        if widths != expected_widths:
            raise RuntimeError(f"Unexpected DOCX cell widths: {widths}")

    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    if "MoonDepSolve v0.3" not in full_text or "MoonLogLens" in full_text:
        raise RuntimeError("DOCX project identity is stale")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_pdf()
    build_docx()
    validate_docx()
    print(f"Generated one-page PDF: {PDF_PATH}")
    print(f"Generated DOCX: {DOCX_PATH}")


if __name__ == "__main__":
    main()
